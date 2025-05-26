from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.utils.translation import gettext_lazy as _
from django.db.models import Count, Sum, Avg, Q, Exists, OuterRef
from django.http import JsonResponse
from .models import *
from .forms import *
from django.db.models import Q
from django.core.paginator import Paginator
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
import logging
logger = logging.getLogger(__name__)
from django.template.loader import get_template
from django.http import HttpResponse
from xhtml2pdf import pisa
from docx import Document
import csv
from django import forms
from .tasks import send_news_notification_task
from io import BytesIO
from django.core.mail import send_mail, EmailMultiAlternatives
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.conf import settings
from datetime import datetime, timedelta

logger = logging.getLogger(__name__)

def admin_required(user):
    return user.is_authenticated and (user.is_superuser or user.role == 'ADMIN')


def get_paginated_queryset(request, queryset, default_per_page=30):
    """Helper function to paginate querysets with customizable per_page."""
    per_page = request.GET.get('per_page', default_per_page)
    try:
        per_page = int(per_page)
        if per_page not in [30, 50, 100]:
            per_page = default_per_page
    except ValueError:
        per_page = default_per_page
    paginator = Paginator(queryset, per_page)
    page = request.GET.get('page')
    paginated = paginator.get_page(page)
    return paginated, per_page



@login_required
def download_reports(request):
    report_type = request.GET.get('type')
    response = None

    if report_type == 'users' and (request.user.is_superuser or request.user.role == 'ADMIN' or request.user.role == 'MAIN_REVEREND'):
        queryset = User.objects.select_related('created_by').prefetch_related('churches')
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.filter(main_reverend=request.user).first()
            queryset = queryset.filter(churches__district=district).distinct() if district else queryset.none()
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="users_list.csv"'
        writer = csv.writer(response)
        writer.writerow([_('First Name'), _('Last Name'), _('Username'), _('Role'), _('Churches')])
        for user in queryset:
            churches = ', '.join([church.name for church in user.churches.all()])
            writer.writerow([user.first_name, user.last_name, user.username, user.get_role_display(), churches])

    elif report_type == 'churches' and (request.user.is_superuser or request.user.role in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']):
        queryset = Church.objects.select_related('district__province', 'pastor')
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.filter(main_reverend=request.user).first()
            queryset = queryset.filter(district=district) if district else queryset.none()
        elif request.user.role == 'PASTOR':
            church = Church.objects.filter(pastor=request.user).first()
            queryset = queryset.filter(id=church.id) if church else queryset.none()
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="churches_list.csv"'
        writer = csv.writer(response)
        writer.writerow([_('Name'), _('District'), _('Province'), _('Pastor')])
        for church in queryset:
            pastor_name = church.pastor.get_full_name() if church.pastor else _('None')
            writer.writerow([church.name, church.district.name, church.district.province.name, pastor_name])

    elif report_type == 'worshipers' and (request.user.is_superuser or request.user.role in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']):
        queryset = Worshiper.objects.select_related('church__district__province')
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.filter(main_reverend=request.user).first()
            queryset = queryset.filter(church__district=district) if district else queryset.none()
        elif request.user.role == 'PASTOR':
            church = Church.objects.filter(pastor=request.user).first()
            queryset = queryset.filter(church=church) if church else queryset.none()
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="worshipers_list.csv"'
        writer = csv.writer(response)
        writer.writerow([_('First Name'), _('Last Name'), _('Church'), _('Gender'), _('Age'), _('Baptized')])
        for worshiper in queryset:
            writer.writerow([
                worshiper.first_name,
                worshiper.last_name,
                worshiper.church.name,
                worshiper.get_gender_display(),
                worshiper.get_age() or _('Unknown'),
                _('Yes') if worshiper.is_baptized else _('No')
            ])

    
    elif report_type == 'deceased_worshipers' and (request.user.is_superuser or request.user.role in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']):
        queryset = Worshiper.objects.select_related('church__district__province').filter(is_deceased=True)
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.filter(main_reverend=request.user).first()
            queryset = queryset.filter(church__district=district) if district else queryset.none()
        elif request.user.role == 'PASTOR':
            church = Church.objects.filter(pastor=request.user).first()
            queryset = queryset.filter(church=church) if church else queryset.none()
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="deceased_worshipers_list.csv"'
        writer = csv.writer(response)
        writer.writerow([_('First Name'), _('Last Name'), _('Church'), _('Gender'), _('Age'), _('Baptized'), _('Deceased')])
        for worshiper in queryset:
            writer.writerow([
                worshiper.first_name,
                worshiper.last_name,
                worshiper.church.name,
                worshiper.get_gender_display(),
                worshiper.get_age() or _('Unknown'),
                _('Yes') if worshiper.is_baptized else _('No'),
                _('Yes'),  # All are deceased
            ])

    


    elif report_type == 'pastors' and (request.user.is_superuser or request.user.role in ['ADMIN', 'MAIN_REVEREND']):
        queryset = User.objects.filter(role='PASTOR').select_related('created_by').prefetch_related('churches')
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.filter(main_reverend=request.user).first()
            queryset = queryset.filter(churches__district=district).distinct() if district else queryset.none()
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="pastors_list.csv"'
        writer = csv.writer(response)
        writer.writerow([_('First Name'), _('Last Name'), _('Username'), _('Gender'), _('Church'),_('Deceased')])
        for pastor in queryset:
            churches = ', '.join([church.name for church in pastor.churches.all()])
            writer.writerow([pastor.first_name, pastor.last_name, pastor.username, pastor.get_gender_display(), churches])

    elif report_type == 'districts' and (request.user.is_superuser or request.user.role in ['ADMIN', 'MAIN_REVEREND']):
        queryset = District.objects.select_related('province', 'main_reverend')
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.filter(main_reverend=request.user).first()
            queryset = queryset.filter(id=district.id) if district else queryset.none()
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="districts_list.csv"'
        writer = csv.writer(response)
        writer.writerow([_('Name'), _('Province'), _('Main Reverend')])
        for district in queryset:
            reverend_name = district.main_reverend.get_full_name() if district.main_reverend else _('None')
            writer.writerow([district.name, district.province.name, reverend_name])

    else:
       context = get_full_report_data(request.user)
       context['user'] = request.user
       logger.info(f"Generating full report PDF with context: {context}")
       template = render_to_string('church/reports/full_report.html', context)
       response = HttpResponse(content_type='application/pdf')
       response['Content-Disposition'] = 'attachment; filename="report.pdf"'
       pisa_status = pisa.CreatePDF(template, dest=response)
       if pisa_status.err:
           logger.error("Error generating PDF report")
           messages.error(request, _('Error generating PDF report.'))
           return redirect('dashboard')

    logger.info(f"Downloaded {report_type} report by {request.user.username} ({request.user.role})")
    return response

@login_required
def download_worshipers(request):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    format = request.GET.get('format', 'pdf')
    
    # Get worshipers based on user role
    if request.user.role == 'PASTOR':
        church = Church.objects.get(pastor=request.user)
        worshipers = Worshiper.objects.filter(church=church)
    elif request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        worshipers = Worshiper.objects.filter(church__district=district)
    else:
        worshipers = Worshiper.objects.all()
    
    if format == 'pdf':
        return generate_pdf(request, 'church/reports/worshipers_pdf.html', {'worshipers': worshipers})
    elif format == 'word':
        return generate_word(request, 'church/reports/worshipers_word.html', {'worshipers': worshipers})
    elif format == 'csv':
        return generate_csv(worshipers, 'worshipers')
    else:
        messages.error(request, _('Invalid format requested'))
        return redirect('worshiper_list')

def generate_pdf(request, template_name, context):
    template = get_template(template_name)
    html = template.render(context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="report.pdf"'
    
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('We had some errors generating the PDF')
    return response

def generate_word(request, template_name, context):
    template = get_template(template_name)
    html = template.render(context)
    
    # Convert HTML to Word - this is a simplified version
    document = Document()
    document.add_heading('Church Management System Report', 0)
    
    # Add content based on context
    if 'worshipers' in context:
        for worshiper in context['worshipers']:
            document.add_paragraph(f"{worshiper.first_name} {worshiper.last_name}")
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=report.docx'
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    return response

def generate_csv(queryset, model_name):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{model_name}_report.csv"'
    
    writer = csv.writer(response)
    
    # Write headers
    if model_name == 'worshipers':
        writer.writerow(['First Name', 'Last Name', 'Gender', 'Baptized', 'Church'])
        for obj in queryset:
            writer.writerow([obj.first_name, obj.last_name, obj.get_gender_display(), 
                           'Yes' if obj.baptized else 'No', obj.church.name])
    
    return response

def get_attendance_data(user):
    if user.role == 'ADMIN':
        attendance_data = AttendanceRecord.objects.values('date').annotate(
            total=Sum('male_count') + Sum('female_count') + Sum('children_count')
        ).order_by('date')[:30]
    elif user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=user)
        attendance_data = AttendanceRecord.objects.filter(
            church__district=district
        ).values('date').annotate(
            total=Sum('male_count') + Sum('female_count') + Sum('children_count')
        ).order_by('date')[:30]
    else:
        attendance_data = []
    
    return {
        'attendance_data': attendance_data,
        'user': user
    }

def get_full_report_data(user):
    if user.role == 'ADMIN':
        return {
            'total_provinces': Province.objects.count(),
            'total_districts': District.objects.count(),
            'total_churches': Church.objects.count(),
            'total_pastors': User.objects.filter(role='PASTOR').count(),
            'total_reverends': User.objects.filter(role='MAIN_REVEREND').count(),
            'total_worshipers': Worshiper.objects.count(),
            'user': user
        }
    elif user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=user)
        return {
            'district': district,
            'total_churches': Church.objects.filter(district=district).count(),
            'total_pastors': User.objects.filter(role='PASTOR', pastor_church__district=district).count(),
            'total_worshipers': Worshiper.objects.filter(church__district=district).count(),
            'user': user
        }
    return {}


@login_required
def mark_worshiper_deceased(request, worshiper_id):
    worshiper = get_object_or_404(Worshiper, id=worshiper_id)
    # Role-based access check
    if request.user.role == 'MAIN_REVEREND':
        district = District.objects.filter(main_reverend=request.user).first()
        if not district or worshiper.church.district != district:
            messages.error(request, _('You do not have permission to modify this worshiper.'))
            return redirect('church:worshiper_list')
    elif request.user.role == 'PASTOR':
        church = Church.objects.filter(pastor=request.user).first()
        if not church or worshiper.church != church:
            messages.error(request, _('You do not have permission to modify this worshiper.'))
            return redirect('church:worshiper_list')
    elif not (request.user.is_superuser or request.user.role == 'ADMIN'):
        messages.error(request, _('You do not have permission to modify this worshiper.'))
        return redirect('church:worshiper_list')

    if request.method == 'POST':
        worshiper.is_deceased = not worshiper.is_deceased  # Toggle deceased status
        worshiper.save()
        status = _('deceased') if worshiper.is_deceased else _('alive')
        messages.success(request, _(f'Worshiper {worshiper} marked as {status}.'))
        logger.info(f"Worshiper {worshiper.id} marked as {status} by {request.user.username} ({request.user.role})")
        return redirect('church:worshiper_list')
    return render(request, 'church/worshiper_deceased_confirm.html', {'worshiper': worshiper})

@login_required
def delete_worshiper(request, worshiper_id):
    worshiper = get_object_or_404(Worshiper, id=worshiper_id)
    # Role-based access check
    if request.user.role == 'MAIN_REVEREND':
        district = District.objects.filter(main_reverend=request.user).first()
        if not district or worshiper.church.district != district:
            messages.error(request, _('You do not have permission to delete this worshiper.'))
            return redirect('church:worshiper_list')
    elif request.user.role == 'PASTOR':
        church = Church.objects.filter(pastor=request.user).first()
        if not church or worshiper.church != church:
            messages.error(request, _('You do not have permission to delete this worshiper.'))
            return redirect('church:worshiper_list')
    elif not (request.user.is_superuser or request.user.role == 'ADMIN'):
        messages.error(request, _('You do not have permission to delete this worshiper.'))
        return redirect('church:worshiper_list')

    if request.method == 'POST':
        worshiper_name = str(worshiper)
        worshiper.delete()
        messages.success(request, _(f'Worshiper {worshiper_name} deleted successfully.'))
        logger.info(f"Worshiper {worshiper_id} deleted by {request.user.username} ({request.user.role})")
        return redirect('church:worshiper_list')
    return render(request, 'church/worshiper_delete_confirm.html', {'worshiper': worshiper})

@login_required
def province_create(request):
    if request.method == 'POST':
        form = ProvinceForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Province created successfully!')
            return redirect('province_list')
    else:
        form = ProvinceForm()
    return render(request, 'church/location/province_form.html', {'form': form})

@login_required
def district_create(request):
    if request.method == 'POST':
        form = DistrictForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'District created successfully!')
            return redirect('district_list')
    else:
        form = DistrictForm()
    return render(request, 'church/location/district_form.html', {'form': form})

@login_required
def church_create(request):
    if not (request.user.is_superuser or request.user.role in ['ADMIN', 'MAIN_REVEREND']):
        messages.error(request, _('Permission denied'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = ChurchForm(request.POST, user=request.user)
        if form.is_valid():
            church = form.save()
            messages.success(request, _('Church created successfully!'))
            return redirect('church_list')
    else:
        form = ChurchForm(user=request.user)
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.get(main_reverend=request.user)
            form.fields['district'].initial = district
            form.fields['district'].widget = forms.HiddenInput()
    
    return render(request, 'church/reverend/church_form.html', {'form': form})
    

@login_required
def profile_view(request):
    if request.method == 'POST':
        form = UserProfileForm(request.POST, request.FILES, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, _('Profile updated successfully!'))
            return redirect('profile')
    else:
        form = UserProfileForm(instance=request.user)
    
    return render(request, 'church/profile.html', {'form': form})

def login_view(request):
    if request.user.is_authenticated:
        return redirect('church:dashboard')
    
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            login(request, user)
            messages.success(request, _('Welcome back, %(name)s!') % {'name': user.get_full_name() or user.username})
            return redirect('church:dashboard')
        else:
            messages.error(request, _('Invalid username or password.'))
    
    return render(request, 'church/auth/login.html')

@login_required
def logout_view(request):
    logout(request)
    messages.success(request, _('You have been logged out successfully.'))
    return redirect('church:login')




@login_required
def dashboard(request):
    if not request.user.is_authenticated:
        messages.info(request, _('Please login to access the dashboard.'))
        return redirect('church:login')
    context = {}
    one_year_ago = datetime.now().date() - timedelta(days=365)
    
    if request.user.is_superuser or request.user.role == 'ADMIN':
        # Admin dashboard
        context['provinces'] = Province.objects.count()
        context['districts'] = District.objects.count()
        context['churches'] = Church.objects.count()
        context['main_reverends'] = User.objects.filter(role='MAIN_REVEREND').count()
        context['pastors'] = User.objects.filter(role='PASTOR').count()
        context['worshipers'] = Worshiper.objects.count()
        context['attendance_data'] = Attendance.objects.values('date').annotate(total=Sum('total')).order_by('date')[:30]
        # Quick Stats
        context['baptized_count'] = Worshiper.objects.filter(is_baptized=True).count()
        context['avg_attendance'] = Attendance.objects.aggregate(avg=Avg('total'))['avg'] or 0
        context['new_members'] = Worshiper.objects.filter(created_at__gte=one_year_ago).count()
        context['gender_distribution'] = Worshiper.objects.values('gender').annotate(count=Count('gender')).order_by('gender')
    
    elif request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        context['district'] = district
        context['churches'] = Church.objects.filter(district=district).count()
        context['pastors'] = User.objects.filter(
            Q(role='PASTOR') & (Q(churches__district=district) | Q(created_by=request.user))
        ).distinct().count()
        context['worshipers'] = Worshiper.objects.filter(church__district=district).count()
        context['attendance_data'] = Attendance.objects.filter(
            church__district=district
        ).values('date').annotate(total=Sum('total')).order_by('date')[:30]
        # Quick Stats
        context['baptized_count'] = Worshiper.objects.filter(church__district=district, is_baptized=True).count()
        context['avg_attendance'] = Attendance.objects.filter(church__district=district).aggregate(avg=Avg('total'))['avg'] or 0
        context['new_members'] = Worshiper.objects.filter(
            church__district=district, created_at__gte=one_year_ago
        ).count()
        context['gender_distribution'] = Worshiper.objects.filter(
            church__district=district
        ).values('gender').annotate(count=Count('gender')).order_by('gender')
    
    elif request.user.role == 'PASTOR':
        church = Church.objects.filter(pastor=request.user).first()
        context['church'] = church
        context['worshipers'] = Worshiper.objects.filter(church=church).count() if church else 0
        context['attendance_data'] = Attendance.objects.filter(
            church=church
        ).values('date').annotate(total=Sum('total')).order_by('date')[:30] if church else []
        # Quick Stats (optional, limited for Pastors)
        context['baptized_count'] = Worshiper.objects.filter(church=church, is_baptized=True).count() if church else 0
        context['avg_attendance'] = Attendance.objects.filter(church=church).aggregate(avg=Avg('total'))['avg'] or 0 if church else 0
        context['new_members'] = Worshiper.objects.filter(
            church=church, created_at__gte=one_year_ago
        ).count() if church else 0
        context['gender_distribution'] = Worshiper.objects.filter(
            church=church
        ).values('gender').annotate(count=Count('gender')).order_by('gender') if church else []
    
    context['news'] = News.objects.order_by('-created_at')[:5]
    logger.info(f"Dashboard for {request.user.username} ({request.user.role}): {context}")
    
    return render(request, 'church/dashboard.html', context)


# Admin-only views
@login_required
def user_list(request):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    queryset = User.objects.all().order_by('role', 'last_name', 'first_name')
    users, per_page = get_paginated_queryset(request, queryset)
    context = {
        'users': users,
        'per_page': per_page,
    }
    return render(request, 'church/admin/user_list.html', context)

@login_required
def user_create(request):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST, request.FILES)
        if form.is_valid():
            user = form.save()
            messages.success(request, _('User created successfully!'))
            return redirect('church:user_list')
    else:
        form = CustomUserCreationForm()
    
    return render(request, 'church/admin/user_create.html', {'form': form})



def user_login(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('church:dashboard')
    else:
        form = AuthenticationForm()
    return render(request, 'church/login.html', {'form': form})


@user_passes_test(admin_required)
def user_edit(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if request.method == 'POST':
        form = UserUpdateForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            messages.success(request, _(f'User {user} updated successfully.'))
            logger.info(f"User {user.username} updated by {request.user.username} (ADMIN)")
            return redirect('church:user_list')
        else:
            messages.error(request, _('Please correct the errors below.'))
    else:
        form = UserUpdateForm(instance=user)
    return render(request, 'church/admin/user_edit.html', {'form': form, 'user': user})

@login_required
def user_delete(request, pk):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    user = get_object_or_404(User, pk=pk)
    if request.method == 'POST':
        user.delete()
        messages.success(request, _('User deleted successfully!'))
        return redirect('church:user_list')
    
    return render(request, 'church/admin/user_delete.html', {'user': user})

# Main Reverend views
@login_required
def pastor_list(request):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND'] and not request.user.is_superuser:
        messages.error(request, _('Permission denied'))
        return redirect('dashboard')
    
    district = None
    if request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        has_district_churches = Exists(
            Church.objects.filter(pastor=OuterRef('pk'), district=district)
        )
        queryset = User.objects.filter(
            Q(role='PASTOR') & (Q(churches__district=district) | Q(created_by=request.user))
        ).annotate(
            has_district_churches=has_district_churches
        ).distinct()
        logger.info(f"Main Reverend {request.user.username} pastor list: {queryset.count()} pastors, usernames: {[p.username for p in queryset]}")
    else:
        queryset = User.objects.filter(role='PASTOR')
        logger.info(f"Admin/Superuser {request.user.username} pastor list: {queryset.count()} pastors")
    
    pastors, per_page = get_paginated_queryset(request, queryset)
    context = {
        'pastors': pastors,
        'district': district,
        'per_page': per_page,
    }
    return render(request, 'church/reverend/pastor_list.html', context)

@user_passes_test(admin_required)
def pastor_assign_church(request, pastor_id=None):
       initial = {}
       if pastor_id:
           try:
               initial['pastor'] = User.objects.get(id=pastor_id, role='PASTOR')
           except User.DoesNotExist:
               messages.warning(request, _('Selected pastor does not exist.'))
       church_id = request.GET.get('church')
       if church_id:
           try:
               initial['church'] = Church.objects.get(id=church_id).id
           except Church.DoesNotExist:
               messages.warning(request, _('Selected church does not exist.'))
       if request.method == 'POST':
           form = PastorAssignChurchForm(request.POST)
           if form.is_valid():
               pastor = form.cleaned_data['pastor']
               church = form.cleaned_data['church']
               # Clear any existing assignment for this pastor
               pastor.churches.clear()
               Church.objects.filter(pastor=pastor).update(pastor=None)
               # Clear existing pastor for this church
               if church.pastor:
                   church.pastor.churches.remove(church)
               # Assign new pastor
               church.pastor = pastor
               church.save()
               pastor.churches.add(church)
               messages.success(request, _(f'Pastor {pastor} assigned to {church}.'))
               logger.info(f"Pastor {pastor.username} assigned to {church.name} by {request.user.username} (ADMIN)")
               return redirect('church:church_list')
           else:
               messages.error(request, _('Please correct the errors below.'))
       else:
           form = PastorAssignChurchForm(initial=initial)
       
       context = {'form': form}
       return render(request, 'church/reverend/pastor_assign_church.html', context)

@user_passes_test(admin_required)
def unassign_pastor(request, church_id):
    church = get_object_or_404(Church, id=church_id)
    if request.method == 'POST':
        pastor = church.pastor
        if pastor:
            church.pastor = None
            church.save()
            pastor.churches.remove()
            messages.success(request, _(f'Pastor {pastor} unassigned from {church}.'))
            logger.info(f"Pastor {pastor.username} unassigned from {church.name} by {request.user.username} (ADMIN)")
        return redirect('church:church_list')
    return render(request, 'church/reverend/pastor_unassign_confirm.html', {'church': church})

@login_required
def pastor_create(request):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND']:
        messages.error(request, _('Permission denied'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = PastorCreationForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            pastor = form.save()
            messages.success(request, _('Pastor created successfully!'))
            return redirect('pastor_list')
    else:
        form = PastorCreationForm(user=request.user)
    
    return render(request, 'church/reverend/pastor_create.html', {'form': form})

@login_required
@user_passes_test(lambda u: u.is_superuser)  # Only main reverend can access
def add_pastor(request):
    if request.method == 'POST':
        form = PastorForm(request.POST, request.FILES)
        if form.is_valid():
            # Create user account
            user = User.objects.create_user(
                username=form.cleaned_data['email'],
                email=form.cleaned_data['email'],
                password=User.objects.make_random_password(),
                first_name=form.cleaned_data['first_name'],
                last_name=form.cleaned_data['last_name']
            )
            
            # Create pastor profile
            pastor = form.save(commit=False)
            pastor.user = user
            pastor.save()
            
            # Send welcome email with password reset link
            # (Implement this based on your email setup)
            
            return redirect('pastor_list')
    else:
        form = PastorForm()
    
    return render(request, 'church/add_pastor.html', {'form': form})


@login_required
def pastor_delete(request, pastor_id):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND'] and not request.user.is_superuser:
        messages.error(request, _('Permission denied'))
        return redirect('pastor_list')
    
    pastor = get_object_or_404(User, pk=pastor_id, role='PASTOR')
    
    # For Main Reverends, ensure the pastor is in their district
    if request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        if not pastor.churches.filter(district=district).exists():
            messages.error(request, _('You can only delete pastors in your district'))
            return redirect('pastor_list')
    
    if request.method == 'POST':
        # Update related churches to set pastor to NULL
        pastor.churches.update(pastor=None)
        pastor.delete()
        messages.success(request, _('Pastor deleted successfully!'))
        return redirect('church:pastor_list')
    
    return render(request, 'church/reverend/pastor_confirm_delete.html', {
        'pastor': pastor,
    })



# Pastor views
@login_required
def worshiper_list(request):
    queryset = Worshiper.objects.select_related('church__district__province').order_by('id')
    form = PastorWorshiperSearchForm(request.GET or None, user=request.user)

    if request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        queryset = queryset.filter(church__district=district)
    elif request.user.role == 'PASTOR':
        church = Church.objects.filter(pastor=request.user).first()
        queryset = queryset.filter(church=church) if church else queryset.none()

    if form.is_valid():
        gender = form.cleaned_data.get('gender')
        name = form.cleaned_data.get('name')
        if gender:
            queryset = queryset.filter(gender=gender)
        if name:
            queryset = queryset.filter(Q(first_name__icontains=name) | Q(last_name__icontains=name))

    worshipers, per_page = get_paginated_queryset(request, queryset)
    context = {
        'worshipers': worshipers,
        'form': form,
        'is_baptized': request.GET.get('is_baptized'),
        'per_page': per_page,
    }
    logger.info(f"Worshiper list for {request.user.username} ({request.user.role}): {context}")
    return render(request, 'church/pastor/worshiper_list.html', context)

@login_required
def worshiper_create(request):
    if request.method == 'POST':
        form = WorshiperForm(request.POST, user=request.user)
        if form.is_valid():
            worshiper = form.save()
            logger.info(f"Worshiper created by {request.user.username}: {worshiper}")
            messages.success(request, _('Worshiper created successfully!'))
            return redirect('church:worshiper_list')
        else:
            logger.error(f"Form errors in worshiper_create: {form.errors}")
    else:
        form = WorshiperForm(user=request.user)  # Line 559
    return render(request, 'church/pastor/worshiper_form.html', {'form': form})

@login_required
def worshiper_edit(request, pk):
    if request.user.role not in ['ADMIN', 'PASTOR']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    worshiper = get_object_or_404(Worshiper, pk=pk)
    
    # Check if pastor is trying to edit a worshiper not in their church
    if request.user.role == 'PASTOR' and worshiper.church.pastor != request.user:
        messages.error(request, _('You can only edit worshipers in your church.'))
        return redirect('worshiper_list')
    
    if request.method == 'POST':
        form = WorshiperForm(request.POST, request.FILES, instance=worshiper)
        if form.is_valid():
            form.save()
            messages.success(request, _('Worshiper updated successfully!'))
            return redirect('worshiper_list')
    else:
        form = WorshiperForm(instance=worshiper)
        if request.user.role == 'PASTOR':
            form.fields['church'].queryset = Church.objects.filter(pastor=request.user)
            form.fields['church'].widget = forms.HiddenInput()
    
    return render(request, 'church/pastor/worshiper_edit.html', {'form': form, 'worshiper': worshiper})

@login_required
def worshiper_delete(request, pk):
    if request.user.role not in ['ADMIN', 'PASTOR']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    worshiper = get_object_or_404(Worshiper, pk=pk)
    
    # Check if pastor is trying to delete a worshiper not in their church
    if request.user.role == 'PASTOR' and worshiper.church.pastor != request.user:
        messages.error(request, _('You can only delete worshipers in your church.'))
        return redirect('worshiper_list')
    
    if request.method == 'POST':
        worshiper.delete()
        messages.success(request, _('Worshiper deleted successfully!'))
        return redirect('worshiper_list')
    
    return render(request, 'church/pastor/worshiper_delete.html', {'worshiper': worshiper})

# Attendance views
@login_required
def attendance_list(request):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    if request.user.role == 'PASTOR':
        church = Church.objects.get(pastor=request.user)
        queryset = AttendanceRecord.objects.filter(church=church)
    elif request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        queryset = AttendanceRecord.objects.filter(church__district=district)
    else:
        queryset = AttendanceRecord.objects.all()
    
    queryset = queryset.order_by('-date')
    page_obj, per_page = get_paginated_queryset(request, queryset)
    context = {
        'page_obj': page_obj,
        'per_page': per_page,
    }
    return render(request, 'church/attendance/list.html', context)

@login_required
def attendance_create(request):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = AttendanceRecordForm(request.POST, user=request.user)
        if form.is_valid():
            record = form.save()
            messages.success(request, _('Attendance record created successfully!'))
            return redirect('attendance_list')
    else:
        form = AttendanceRecordForm(user=request.user)
    
    return render(request, 'church/attendance/create.html', {'form': form})

@login_required
def attendance_edit(request, pk):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    record = get_object_or_404(AttendanceRecord, pk=pk)
    
    # Check permissions
    if request.user.role == 'PASTOR' and record.church.pastor != request.user:
        messages.error(request, _('You can only edit records from your church.'))
        return redirect('attendance_list')
    elif request.user.role == 'MAIN_REVEREND' and record.church.district.main_reverend != request.user:
        messages.error(request, _('You can only edit records from your district.'))
        return redirect('attendance_list')
    
    if request.method == 'POST':
        form = AttendanceRecordForm(request.POST, instance=record, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, _('Attendance record updated successfully!'))
            return redirect('attendance_list')
    else:
        form = AttendanceRecordForm(instance=record, user=request.user)
    
    return render(request, 'church /attendance/edit.html',{'form':form ,'record':record})

@login_required
def attendance_delete(request, pk):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND', 'PASTOR']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    record = get_object_or_404(AttendanceRecord, pk=pk)
    
    # Check permissions
    if request.user.role == 'PASTOR' and record.church.pastor != request.user:
        messages.error(request, _('You can only delete records from your church.'))
        return redirect('attendance_list')
    elif request.user.role == 'MAIN_REVEREND' and record.church.district.main_reverend != request.user:
        messages.error(request, _('You can only delete records from your district.'))
        return redirect('attendance_list')
    
    if request.method == 'POST':
        record.delete()
        messages.success(request, _('Attendance record deleted successfully!'))
        return redirect('attendance_list')
    
    return render(request, 'church/attendance/delete.html', {'record': record})

# News views
@login_required
def news_list(request):
    if request.user.is_superuser or request.user.role == 'ADMIN':
        queryset = News.objects.all()
    elif request.user.role == 'MAIN_REVEREND':
        queryset = News.objects.filter(
            is_published=True,
            target_audience__in=['ALL', 'REVERENDS']
        )
    elif request.user.role == 'PASTOR':
        queryset = News.objects.filter(
            is_published=True,
            target_audience__in=['ALL', 'REVERENDS', 'PASTORS']
        )
    else:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    queryset = queryset.order_by('-created_at')
    page_obj, per_page = get_paginated_queryset(request, queryset)
    context = {
        'page_obj': page_obj,
        'per_page': per_page,
    }
    return render(request, 'church/news/list.html', context)

@login_required
def news_create(request):
    if request.method == 'POST':
        form = NewsForm(request.POST, user=request.user)
        if form.is_valid():
            news = form.save(commit=False)
            news.created_by = request.user
            news.save()
            # Get worshipers for email notifications
            queryset = Worshiper.objects.select_related('church__district__province').order_by('id')
            if request.user.role == 'MAIN_REVEREND':
                district = District.objects.get(main_reverend=request.user)
                queryset = queryset.filter(church__district=district)
            elif request.user.role == 'PASTOR':
                church = Church.objects.filter(pastor=request.user).first()
                queryset = queryset.filter(church=church) if church else queryset.none()
            
            # Paginate worshiper list for display or batch processing
            paginator = Paginator(queryset, 10)  # Line 534
            page_number = request.GET.get('page')
            worshipers = paginator.get_page(page_number)
            
            # Send email notifications
            subject = news.title
            html_message = render_to_string('church/news/email_notification.html', {'news': news})
            plain_message = strip_tags(html_message)
            from_email = settings.EMAIL_HOST_USER
            recipient_list = [w.email for w in queryset if w.email]
            
            try:
                email = EmailMultiAlternatives(
                    subject, plain_message, from_email, recipient_list
                )
                email.attach_alternative(html_message, "text/html")
                email.send()
                messages.success(request, _('News created and emails sent successfully!'))
            except Exception as e:
                logger.error(f"Error sending news emails: {str(e)}")
                messages.error(request, _('News created, but failed to send emails.'))
            
            return redirect('church:news_list')
    else:
        form = NewsForm(user=request.user)
    
    # Paginate worshiper list for recipient selection
    queryset = Worshiper.objects.select_related('church__district__province').order_by('id')
    if request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        queryset = queryset.filter(church__district=district)
    elif request.user.role == 'PASTOR':
        church = Church.objects.filter(pastor=request.user).first()
        queryset = queryset.filter(church=church) if church else queryset.none()
    
    paginator = Paginator(queryset, 10)
    page_number = request.GET.get('page')
    worshipers = paginator.get_page(page_number)
    
    context = {
        'form': form,
        'worshipers': worshipers,
    }
    return render(request, 'church/news/create.html', context)

@login_required
def news_edit(request, pk):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    news = get_object_or_404(News, pk=pk)
    
    if news.author != request.user and not (request.user.is_superuser or request.user.role == 'ADMIN'):
        messages.error(request, _('You can only edit your own news items.'))
        return redirect('news_list')
    
    if request.method == 'POST':
        form = NewsForm(request.POST, instance=news, user=request.user)
        if form.is_valid():
            news = form.save()
            
            # Send email notifications if newly published or changed and requested
            if (news.is_published and form.cleaned_data.get('send_notification', True) and 
                (not news.emails_sent or form.has_changed())):
                send_news_notification_task.delay(news.pk)  # Correct way to call the task
                
            messages.success(request, _('News item updated successfully!'))
            return redirect('news_list')
    else:
        form = NewsForm(instance=news, user=request.user)
    
    return render(request, 'church/news/edit.html', {'form': form, 'news': news})

def send_news_notification(news):
    # Determine which users should receive this news
    if news.target_audience == 'ALL':
        recipients = User.objects.filter(is_active=True)
    elif news.target_audience == 'REVERENDS':
        recipients = User.objects.filter(
            Q(role='MAIN_REVEREND') | Q(role='ADMIN') | Q(is_superuser=True),
            is_active=True
        )
    elif news.target_audience == 'PASTORS':
        recipients = User.objects.filter(
            Q(role='PASTOR') | Q(role='MAIN_REVEREND') | Q(role='ADMIN') | Q(is_superuser=True),
            is_active=True
        )
    else:
        recipients = User.objects.none()
    
    # Prepare email content
    subject = f"New Announcement: {news.title}"
    html_content = render_to_string('church/news/email_notification.html', {
        'news': news,
        'site_name': 'Church Management System'
    })
    text_content = strip_tags(html_content)
    
    # Send emails (in production, consider using Celery for this)
    for user in recipients:
        if user.email:  # Only send to users with email addresses
            try:
                msg = EmailMultiAlternatives(
                    subject,
                    text_content,
                    settings.DEFAULT_FROM_EMAIL,
                    [user.email]
                )
                msg.attach_alternative(html_content, "text/html")
                msg.send()
            except Exception as e:
                logger.error(f"Failed to send news email to {user.email}: {str(e)}")
    
    # Update news record to indicate emails were sent
    news.emails_sent = True
    news.email_sent_at = timezone.now()
    news.save()


@login_required
def news_delete(request, pk):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND']:
        messages.error(request, _('You do not have permission to access this page.'))
        return redirect('dashboard')
    
    news = get_object_or_404(News, pk=pk)
    
    # Check if user is the author or admin
    if news.author != request.user and not (request.user.is_superuser or request.user.role == 'ADMIN'):
        messages.error(request, _('You can only delete your own news items.'))
        return redirect('news_list')
    
    if request.method == 'POST':
        news.delete()
        messages.success(request, _('News item deleted successfully!'))
        return redirect('church:news_list')
    
    return render(request, 'church/news/delete.html', {'news': news})

@login_required
def news_detail(request, pk):
    news = get_object_or_404(News, pk=pk)
    
    # Check if user has permission to view this news
    if news.target_audience == 'REVERENDS' and request.user.role not in ['ADMIN', 'MAIN_REVEREND']:
        messages.error(request, _('You do not have permission to view this news item.'))
        return redirect('church:dashboard')
    elif news.target_audience == 'PASTORS' and request.user.role != 'PASTOR':
        messages.error(request, _('You do not have permission to view this news item.'))
        return redirect('dashboard')
    
    return render(request, 'church/news/detail.html', {
        'news': news,
        'user': request.user
    })

# Reports and Analytics
@login_required
def dashboard(request):
    context = {}
    one_year_ago = datetime.now().date() - timedelta(days=365)
    
    if request.user.is_superuser or request.user.role == 'ADMIN':
        context['provinces'] = Province.objects.count()
        context['districts'] = District.objects.count()
        context['churches'] = Church.objects.count()
        context['main_reverends'] = User.objects.filter(role='MAIN_REVEREND').count()
        context['pastors'] = User.objects.filter(role='PASTOR').count()
        context['worshipers'] = Worshiper.objects.count()
        context['attendance_data'] = Attendance.objects.values('date').annotate(total=Sum('total')).order_by('date')[:30]
        context['baptized_count'] = Worshiper.objects.filter(is_baptized=True).count()
        context['avg_attendance'] = Attendance.objects.aggregate(avg=Avg('total'))['avg'] or 0
        context['new_members'] = Worshiper.objects.filter(created_at__gte=one_year_ago).count()
        context['gender_distribution'] = Worshiper.objects.values('gender').annotate(count=Count('gender')).order_by('gender')
    
    elif request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        context['district'] = district
        context['churches'] = Church.objects.filter(district=district).count()
        context['pastors'] = User.objects.filter(
            Q(role='PASTOR') & (Q(churches__district=district) | Q(created_by=request.user))
        ).distinct().count()
        context['worshipers'] = Worshiper.objects.filter(church__district=district).count()
        context['attendance_data'] = Attendance.objects.filter(
            church__district=district
        ).values('date').annotate(total=Sum('total')).order_by('date')[:30]
        context['baptized_count'] = Worshiper.objects.filter(church__district=district, is_baptized=True).count()
        context['avg_attendance'] = Attendance.objects.filter(church__district=district).aggregate(avg=Avg('total'))['avg'] or 0
        context['new_members'] = Worshiper.objects.filter(
            church__district=district, created_at__gte=one_year_ago
        ).count()
        context['gender_distribution'] = Worshiper.objects.filter(
            church__district=district
        ).values('gender').annotate(count=Count('gender')).order_by('gender')
    
    elif request.user.role == 'PASTOR':
        church = Church.objects.filter(pastor=request.user).first()
        context['church'] = church
        context['worshipers'] = Worshiper.objects.filter(church=church).count() if church else 0
        context['attendance_data'] = Attendance.objects.filter(
            church=church
        ).values('date').annotate(total=Sum('total')).order_by('date')[:30] if church else []
        context['baptized_count'] = Worshiper.objects.filter(church=church, is_baptized=True).count() if church else 0
        context['avg_attendance'] = Attendance.objects.filter(church=church).aggregate(avg=Avg('total'))['avg'] or 0 if church else 0
        context['new_members'] = Worshiper.objects.filter(
            church=church, created_at__gte=one_year_ago
        ).count() if church else 0
        context['gender_distribution'] = Worshiper.objects.filter(
            church=church
        ).values('gender').annotate(count=Count('gender')).order_by('gender') if church else []
    
    context['news'] = News.objects.order_by('-created_at')[:5]
    logger.info(f"Dashboard for {request.user.username} ({request.user.role}): {context}")
    
    return render(request, 'church/dashboard.html', context)

@login_required
def reports_dashboard(request):
    context = {}
    one_year_ago = datetime.now().date() - timedelta(days=365)
    
    if request.user.is_superuser or request.user.role == 'ADMIN':
        context['total_worshipers'] = Worshiper.objects.count()
        context['baptized_count'] = Worshiper.objects.filter(is_baptized=True).count()
        context['avg_attendance'] = Attendance.objects.aggregate(avg=Avg('total'))['avg'] or 0
        context['new_members'] = Worshiper.objects.filter(created_at__gte=one_year_ago).count()
        context['gender_distribution'] = Worshiper.objects.values('gender').annotate(count=Count('gender')).order_by('gender')
        context['attendance_data'] = Attendance.objects.values('date').annotate(total=Sum('total')).order_by('date')[:90]
        context['average_age'] = Worshiper.objects.filter(date_of_birth__isnull=False).aggregate(
            avg_age=Avg(datetime.now().date().year - models.F('date_of_birth__year'))
        )['avg_age'] or 0
    
    elif request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        context['district'] = district
        context['total_worshipers'] = Worshiper.objects.filter(church__district=district).count()
        context['baptized_count'] = Worshiper.objects.filter(church__district=district, is_baptized=True).count()
        context['avg_attendance'] = Attendance.objects.filter(church__district=district).aggregate(avg=Avg('total'))['avg'] or 0
        context['new_members'] = Worshiper.objects.filter(church__district=district, created_at__gte=one_year_ago).count()
        context['gender_distribution'] = Worshiper.objects.filter(
            church__district=district
        ).values('gender').annotate(count=Count('gender')).order_by('gender')
        context['attendance_data'] = Attendance.objects.filter(
            church__district=district
        ).values('date').annotate(total=Sum('total')).order_by('date')[:90]
        context['average_age'] = Worshiper.objects.filter(
            church__district=district, date_of_birth__isnull=False
        ).aggregate(
            avg_age=Avg(datetime.now().date().year - models.F('date_of_birth__year'))
        )['avg_age'] or 0
    
    elif request.user.role == 'PASTOR':
        church = Church.objects.filter(pastor=request.user).first()
        context['church'] = church
        context['total_worshipers'] = Worshiper.objects.filter(church=church).count() if church else 0
        context['baptized_count'] = Worshiper.objects.filter(church=church, is_baptized=True).count() if church else 0
        context['avg_attendance'] = Attendance.objects.filter(church=church).aggregate(avg=Avg('total'))['avg'] or 0 if church else 0
        context['new_members'] = Worshiper.objects.filter(
            church=church, created_at__gte=one_year_ago
        ).count() if church else 0
        context['gender_distribution'] = Worshiper.objects.filter(
            church=church
        ).values('gender').annotate(count=Count('gender')).order_by('gender') if church else []
        context['attendance_data'] = Attendance.objects.filter(
            church=church
        ).values('date').annotate(total=Sum('total')).order_by('date')[:90] if church else []
        context['average_age'] = Worshiper.objects.filter(
            church=church, date_of_birth__isnull=False
        ).aggregate(
            avg_age=Avg(datetime.now().date().year - models.F('date_of_birth__year'))
        )['avg_age'] or 0 if church else 0
    
    logger.info(f"Reports dashboard for {request.user.username} ({request.user.role}): {context}")
    
    return render(request, 'church/reports.html', context)

# API endpoints for charts
@login_required
def attendance_chart_data(request):
    if request.user.role == 'ADMIN':
        data = AttendanceRecord.objects.values('date').annotate(
            total=Sum('male_count') + Sum('female_count') + Sum('children_count')
        ).order_by('date')
    elif request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        data = AttendanceRecord.objects.filter(
            church__district=district
        ).values('date').annotate(
            total=Sum('male_count') + Sum('female_count') + Sum('children_count')
        ).order_by('date')
    elif request.user.role == 'PASTOR':
        church = Church.objects.get(pastor=request.user)
        data = AttendanceRecord.objects.filter(
            church=church
        ).values('date').annotate(
            total=Sum('male_count') + Sum('female_count') + Sum('children_count')
        ).order_by('date')
    else:
        return JsonResponse({'error': 'Unauthorized'}, status=403)
    
    labels = [item['date'].strftime('%Y-%m-%d') for item in data]
    values = [item['total'] for item in data]
    
    return JsonResponse({
        'labels': labels,
        'values': values,
    })

# Language switching
def set_language(request):
    if request.method == 'POST':
        language = request.POST.get('language')
        if language in [lang[0] for lang in settings.LANGUAGES]:  # Use settings.LANGUAGES
            if hasattr(request, 'session'):
                request.session['django_language'] = language
            response = redirect(request.META.get('HTTP_REFERER', '/'))
            response.set_cookie(settings.LANGUAGE_COOKIE_NAME, language)
            return response
    return redirect('/')



# views.py - Add these new views

@login_required
def main_reverend_create(request):
    if not (request.user.is_superuser or request.user.role == 'ADMIN'):
        messages.error(request, _('Permission denied'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = MainReverendCreationForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            main_reverend = form.save()
            messages.success(request, _('Main Reverend created successfully!'))
            return redirect('user_list')
    else:
        form = MainReverendCreationForm(user=request.user)
    
    return render(request, 'church/main_reverend_create.html', {'form': form})

@login_required
def pastor_create(request):
    if request.user.role not in ['ADMIN', 'MAIN_REVEREND']:
        messages.error(request, _('Permission denied'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = PastorCreationForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            pastor = form.save()
            messages.success(request, _('Pastor created successfully!'))
            return redirect('church:pastor_list')
    else:
        form = PastorCreationForm(user=request.user)
    
    return render(request, 'church/reverend/pastor_create.html', {'form': form})


@login_required
def province_list(request):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to view this page.'))
        return redirect('dashboard')
    
    queryset = Province.objects.all().order_by('name')
    provinces, per_page = get_paginated_queryset(request, queryset)
    context = {
        'provinces': provinces,
        'per_page': per_page,
    }
    return render(request, 'church/location/province_list.html', context)

@login_required
def province_create(request):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = ProvinceForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, _('Province created successfully!'))
            return redirect('church:province_list')
    else:
        form = ProvinceForm()
    
    return render(request, 'church/location/province_form.html', {'form': form})

@login_required
def province_edit(request, pk):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    province = get_object_or_404(Province, pk=pk)
    if request.method == 'POST':
        form = ProvinceForm(request.POST, instance=province)
        if form.is_valid():
            form.save()
            messages.success(request, _('Province updated successfully!'))
            return redirect('church:province_list')
    else:
        form = ProvinceForm(instance=province)
    
    return render(request, 'church/location/province_form.html', {'form': form})

@login_required
def province_delete(request, pk):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    province = get_object_or_404(Province, pk=pk)
    if request.method == 'POST':
        province.delete()
        messages.success(request, _('Province deleted successfully!'))
        return redirect('church:province_list')
    
    return render(request, 'church/location/province_delete.html', {'province': province})

# District Views
@login_required
def district_list(request):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to view this page.'))
        return redirect('dashboard')
    
    queryset = District.objects.all().order_by('province', 'name')
    districts, per_page = get_paginated_queryset(request, queryset)
    context = {
        'districts': districts,
        'per_page': per_page,
    }
    return render(request, 'church/location/district_list.html', context)

@login_required
def district_create(request):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = DistrictForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, _('District created successfully!'))
            return redirect('church:district_list')
    else:
        form = DistrictForm()
    
    return render(request, 'church/location/district_form.html', {'form': form})

@login_required
def district_edit(request, pk):
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    district = get_object_or_404(District, pk=pk)
    if request.method == 'POST':
        form = DistrictForm(request.POST, instance=district)
        if form.is_valid():
            form.save()
            messages.success(request, _('District updated successfully!'))
            return redirect('church:district_list')
    else:
        form = DistrictForm(instance=district)
    
    return render(request, 'church/location/district_form.html', {'form': form})

@login_required
def district_delete(request, pk):

    district = get_object_or_404(District, pk=pk)
    
    if district.churches.exists():
        messages.error(request, _('Cannot delete district with existing churches.'))
        return redirect('district_list')

    
    if not request.user.is_superuser and request.user.role != 'ADMIN':
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    district = get_object_or_404(District, pk=pk)
    
    if request.method == 'POST':
        district.delete()
        messages.success(request, _('District deleted successfully!'))
        return redirect('church:district_list')
    
    return render(request, 'church/location/district_delete.html', {'district': district})

# Church Views
@login_required
def church_list(request):
    if request.user.is_superuser or request.user.role == 'ADMIN':
        queryset = Church.objects.all().order_by('district', 'name')
    elif request.user.role == 'MAIN_REVEREND':
        district = District.objects.get(main_reverend=request.user)
        queryset = Church.objects.filter(district=district).order_by('name')
    else:
        messages.error(request, _('You do not have permission to view this page.'))
        return redirect('dashboard')
    
    churches, per_page = get_paginated_queryset(request, queryset)
    context = {
        'churches': churches,
        'per_page': per_page,
    }
    return render(request, 'church/reverend/church_list.html', context)

@login_required
def church_create(request):
    if not (request.user.is_superuser or request.user.role in ['ADMIN', 'MAIN_REVEREND']):
        messages.error(request, _('Permission denied'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = ChurchForm(request.POST, user=request.user)
        if form.is_valid():
            church = form.save()
            messages.success(request, _('Church created successfully!'))
            return redirect('church_list')
    else:
        form = ChurchForm(user=request.user)
        if request.user.role == 'MAIN_REVEREND':
            district = District.objects.get(main_reverend=request.user)
            form.fields['district'].initial = district
            form.fields['district'].widget = forms.HiddenInput()
    
    return render(request, 'church/reverend/church_form.html', {'form': form})

@login_required
def assign_church(request, church_id=None):
    if request.user.role != 'MAIN_REVEREND':
        messages.error(request, _('Permission denied'))
        return redirect('church:dashboard')

    district = District.objects.get(main_reverend=request.user)

    if church_id:
        church = get_object_or_404(Church, pk=church_id, district=district)
    else:
        church = None

    available_pastors = User.objects.filter(
        role='PASTOR',
        churches__district=district
    ).distinct()

    class AssignChurchForm(forms.Form):
        pastor = forms.ModelChoiceField(
            queryset=available_pastors,
            required=True,
            label=_('Select Pastor'),
            widget=forms.Select(attrs={'class': 'form-control'})
        )
        church = forms.ModelChoiceField(
            queryset=Church.objects.filter(district=district, pastor__isnull=True),
            required=True,
            label=_('Select Church'),
            widget=forms.Select(attrs={'class': 'form-control'})
        )

    if request.method == 'POST':
        form = AssignChurchForm(request.POST)
        if form.is_valid():
            selected_church = form.cleaned_data['church']
            pastor = form.cleaned_data['pastor']
            selected_church.pastor = pastor
            selected_church.save()
            messages.success(request, _('Church assigned successfully!'))
            return redirect('pastor_list')
    else:
        form = AssignChurchForm(initial={'church': church} if church else {})

    return render(request, 'church/reverend/assign_church.html', {
        'form': form,
        'district': district,
    })

@login_required
def church_edit(request, pk):
    church = get_object_or_404(Church, pk=pk)
    
    if request.user.is_superuser or request.user.role == 'ADMIN':
        form = ChurchForm(instance=church)
    elif request.user.role == 'MAIN_REVEREND':
        if church.district.main_reverend != request.user:
            messages.error(request, _('You can only edit churches in your district.'))
            return redirect('church_list')
        form = ChurchForm(instance=church)
        form.fields['district'].queryset = District.objects.filter(main_reverend=request.user)
    else:
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = ChurchForm(request.POST, instance=church)
        if form.is_valid():
            form.save()
            messages.success(request, _('Church updated successfully!'))
            return redirect('church_list')
    
    return render(request, 'church/location/church_form.html', {'form': form})

@login_required
def church_delete(request, pk):
    church = get_object_or_404(Church, pk=pk)
    
    if request.user.is_superuser or request.user.role == 'ADMIN':
        pass  # Allow deletion
    elif request.user.role == 'MAIN_REVEREND':
        if church.district.main_reverend != request.user:
            messages.error(request, _('You can only delete churches in your district.'))
            return redirect('church_list')
    else:
        messages.error(request, _('You do not have permission to perform this action.'))
        return redirect('dashboard')
    
    if request.method == 'POST':
        church.delete()
        messages.success(request, _('Church deleted successfully!'))
        return redirect('church:church_list')
    
    return render(request, 'church/location/church_delete_confirm.html', {'church': church})




def district_list_api(request):
    province_id = request.GET.get('province_id')
    if province_id:
        districts = District.objects.filter(province_id=province_id).values('id', 'name')
        return JsonResponse({'districts': list(districts)})
    return JsonResponse({'districts': []})



@login_required
@user_passes_test(admin_required)
def search(request):
    user_form = AdminUserSearchForm(request.GET or None)
    church_form = AdminChurchSearchForm(request.GET or None)
    worshiper_form = AdminWorshiperSearchForm(request.GET or None)

    users = User.objects.select_related('created_by').prefetch_related('churches')
    churches = Church.objects.select_related('district__province', 'pastor')
    worshipers = Worshiper.objects.select_related('church__district__province')

    if user_form.is_valid():
        province = user_form.cleaned_data.get('province')
        district = user_form.cleaned_data.get('district')
        church = user_form.cleaned_data.get('church')
        name = user_form.cleaned_data.get('name')
        if province:
            users = users.filter(churches__district__province=province)
        if district:
            users = users.filter(churches__district=district)
        if church:
            users = users.filter(churches=church)
        if name:
            users = users.filter(Q(first_name__icontains=name) | Q(last_name__icontains=name))

    if church_form.is_valid():
        province = church_form.cleaned_data.get('province')
        district = church_form.cleaned_data.get('district')
        name = church_form.cleaned_data.get('name')
        if province:
            churches = churches.filter(district__province=province)
        if district:
            churches = churches.filter(district=district)
        if name:
            churches = churches.filter(name__icontains=name)

    if worshiper_form.is_valid():
        gender = worshiper_form.cleaned_data.get('gender')
        name = worshiper_form.cleaned_data.get('name')
        if gender:
            worshipers = worshipers.filter(gender=gender)
        if name:
            worshipers = worshipers.filter(Q(first_name__icontains=name) | Q(last_name__icontains=name))

    users, user_per_page = get_paginated_queryset(request, users.order_by('id'), default_per_page=30)
    churches, church_per_page = get_paginated_queryset(request, churches.order_by('id'), default_per_page=30)
    worshipers, worshiper_per_page = get_paginated_queryset(request, worshipers.order_by('id'), default_per_page=30)

    context = {
        'user_form': user_form,
        'church_form': church_form,
        'worshiper_form': worshiper_form,
        'users': users,
        'churches': churches,
        'worshipers': worshipers,
        'user_per_page': user_per_page,
        'church_per_page': church_per_page,
        'worshiper_per_page': worshiper_per_page,
    }
    logger.info(f"Admin search by {request.user.username}: {context}")
    return render(request, 'church/search.html', context)

@login_required
@user_passes_test(lambda u: u.is_authenticated and u.role == 'MAIN_REVEREND')
def reverend_search(request):
    church_form = ReverendChurchSearchForm(request.GET or None, user=request.user)
    pastor_form = ReverendPastorSearchForm(request.GET or None, user=request.user)

    churches = Church.objects.select_related('district__province', 'pastor')
    pastors = User.objects.filter(role='PASTOR').select_related('created_by')

    if church_form.is_valid() and church_form.district:
        churches = churches.filter(district=church_form.district)
        name = church_form.cleaned_data.get('name')
        if name:
            churches = churches.filter(name__icontains=name)

    if pastor_form.is_valid():
        church = pastor_form.cleaned_data.get('church')
        gender = pastor_form.cleaned_data.get('gender')
        name = pastor_form.cleaned_data.get('name')
        if church:
            pastors = pastors.filter(churches=church)
        if gender:
            pastors = pastors.filter(gender=gender)
        if name:
            pastors = pastors.filter(Q(first_name__icontains=name) | Q(last_name__icontains=name))

    church_paginator = Paginator(churches.order_by('id'), 10)
    pastor_paginator = Paginator(pastors.order_by('id'), 10)

    church_page = church_paginator.get_page(request.GET.get('church_page'))
    pastor_page = pastor_paginator.get_page(request.GET.get('pastor_page'))

    context = {
        'church_form': church_form,
        'pastor_form': pastor_form,
        'churches': church_page,
        'pastors': pastor_page,
    }
    logger.info(f"Reverend search by {request.user.username}: {context}")
    return render(request, 'church/reverend_search.html', context)


@user_passes_test(admin_required)
def user_create(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            messages.success(request, _(f'User {user} created successfully.'))
            logger.info(f"User {user.username} created by {request.user.username} (ADMIN)")
            return redirect('church:user_list')
        else:
            messages.error(request, _('Please correct the errors below.'))
    else:
        form = UserCreationForm()
    return render(request, 'church/admin/user_create.html', {'form': form})